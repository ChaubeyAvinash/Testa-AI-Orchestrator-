"""
TESTA – Concise Solution Summary Generator (1-2 pages)
Accenture theme: purple #a100ff / dark #7500c0
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Palette ────────────────────────────────────────────────────────────────────
PURPLE      = RGBColor(0xA1, 0x00, 0xFF)
PURPLE_DARK = RGBColor(0x75, 0x00, 0xC0)
BLACK       = RGBColor(0x00, 0x00, 0x00)
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
GREY_DARK   = RGBColor(0x44, 0x44, 0x44)
SUCCESS     = RGBColor(0x00, 0x99, 0x44)
AZURE       = RGBColor(0x00, 0x78, 0xD4)


# ── Helpers ────────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    hex_color = hex_color.lstrip('#')
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color.upper())
    tcPr.append(shd)


def set_row_height(row, height_cm):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(int(height_cm * 360000)))
    trHeight.set(qn('w:hRule'), 'exact')
    trPr.append(trHeight)


def para_spacing(para, before=0, after=0, line=None):
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), str(before))
    spacing.set(qn('w:after'),  str(after))
    if line:
        spacing.set(qn('w:line'),     str(line))
        spacing.set(qn('w:lineRule'), 'auto')
    pPr.append(spacing)


def add_hr(doc, color='A100FF', thickness='6'):
    p = doc.add_paragraph()
    para_spacing(p, 0, 0)
    pPr = p._p.get_or_add_pPr()
    pb = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    thickness)
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), color)
    pb.append(bot)
    pPr.append(pb)
    return p


def section_label(doc, text):
    p = doc.add_paragraph()
    para_spacing(p, before=140, after=40)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = PURPLE
    run.font.name = 'Arial'
    return p


# ── Header ─────────────────────────────────────────────────────────────────────

def make_header(doc):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.rows[0].cells[0]
    set_cell_bg(cell, 'A100FF')
    cp = cell.paragraphs[0]
    cp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = cp.add_run('  >  TESTA  —  SOLUTION SUMMARY')
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = WHITE
    run.font.name = 'Arial'
    para_spacing(cp, before=120, after=120)
    set_row_height(tbl.rows[0], 1.2)

    p = doc.add_paragraph()
    para_spacing(p, before=60, after=20)
    r1 = p.add_run('AI-Powered Autonomous Web Testing Platform   ')
    r1.bold = True
    r1.font.size = Pt(12)
    r1.font.color.rgb = PURPLE
    r1.font.name = 'Arial'
    r2 = p.add_run('"No test scripts. No configuration. Just a URL."')
    r2.italic = True
    r2.font.size = Pt(10.5)
    r2.font.color.rgb = GREY_DARK
    r2.font.name = 'Arial'

    add_hr(doc, 'A100FF', '8')


# ── Section 1: Problem Statement ───────────────────────────────────────────────

def make_problem(doc):
    section_label(doc, '01  Problem Statement')

    problems = [
        'Manual test authoring requires specialist knowledge and weeks of setup per project',
        'Test scripts grow stale — diverging from live application behaviour, creating false CI confidence',
        'No immediate root-cause visibility when CI pipelines fail; engineers spend hours reading stack traces',
        'Junior engineers cannot contribute meaningful test coverage — high skill barrier blocks participation',
        'QA bottlenecks force coverage cuts under delivery pressure, increasing production defect risk',
    ]

    tbl = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    lc = tbl.rows[0].cells[0]
    rc = tbl.rows[0].cells[1]
    lc.width = Inches(1.1)
    rc.width = Inches(5.5)
    set_cell_bg(lc, '111111')
    set_cell_bg(rc, 'FAFAFA')
    lp = lc.paragraphs[0]
    para_spacing(lp, 80, 80)
    lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lr = lp.add_run('PAIN\nPOINTS')
    lr.bold = True
    lr.font.size = Pt(9)
    lr.font.color.rgb = PURPLE
    lr.font.name = 'Arial'

    for i, prob in enumerate(problems):
        rp = rc.paragraphs[0] if i == 0 else rc.add_paragraph()
        para_spacing(rp, 40, 40)
        rr = rp.add_run(f'•  {prob}')
        rr.font.size = Pt(9.5)
        rr.font.name = 'Arial'
        rr.font.color.rgb = BLACK


# ── Section 2: Solution Approach ───────────────────────────────────────────────

def make_solution(doc):
    section_label(doc, '02  Solution Approach')

    intro = doc.add_paragraph()
    para_spacing(intro, before=20, after=50)
    r = intro.add_run(
        'TESTA is a fully autonomous AI testing agent. Provide a URL — the entire '
        'pipeline runs without human intervention, streaming live progress to the browser.'
    )
    r.font.size = Pt(9.5)
    r.font.name = 'Arial'
    r.font.color.rgb = BLACK

    stages = [
        ('01\nCRAWL',    'Playwright BFS\nDiscovers pages,\nforms & inputs',    'A100FF'),
        ('02\nGENERATE', 'Azure GPT-5.1\nWrites complete\n.spec.ts file',       '7500C0'),
        ('03\nEXECUTE',  'Playwright Runner\nRuns all tests &\ncaptures screenshots', 'A100FF'),
        ('04\nANALYZE',  'GPT-5.1 Diagnosis\nPlain-English fix\nper failure',   '7500C0'),
    ]

    tbl = doc.add_table(rows=1, cols=7)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    col_widths = [1.35, 0.18, 1.35, 0.18, 1.35, 0.18, 1.35]
    for ci, w in enumerate(col_widths):
        for row in tbl.rows:
            row.cells[ci].width = Inches(w)

    for si, (name, detail, color) in enumerate(stages):
        ci = si * 2
        cell = tbl.rows[0].cells[ci]
        set_cell_bg(cell, '0d0d0d')

        p_hdr = cell.paragraphs[0]
        p_hdr.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para_spacing(p_hdr, 60, 30)
        r_hdr = p_hdr.add_run(name)
        r_hdr.bold = True
        r_hdr.font.size = Pt(10)
        r_hdr.font.color.rgb = RGBColor.from_string(color)
        r_hdr.font.name = 'Arial'

        p_det = cell.add_paragraph()
        p_det.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para_spacing(p_det, 10, 60)
        r_det = p_det.add_run(detail)
        r_det.font.size = Pt(8.5)
        r_det.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
        r_det.font.name = 'Arial'

        if si < 3:
            arrow_cell = tbl.rows[0].cells[ci + 1]
            set_cell_bg(arrow_cell, '000000')
            ap = arrow_cell.paragraphs[0]
            ap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para_spacing(ap, 120, 0)
            ar = ap.add_run('→')
            ar.font.size = Pt(14)
            ar.font.color.rgb = PURPLE
            ar.font.name = 'Arial'


# ── Section 3: Tech Stack ──────────────────────────────────────────────────────

def make_tech_stack(doc):
    section_label(doc, '03  Tech Stack')

    rows = [
        ('Frontend',       'Next.js 16 (App Router)  ·  React 19  ·  TypeScript  ·  Tailwind CSS v4  ·  Recharts'),
        ('Backend',        'NestJS 11  ·  TypeScript  ·  Prisma v7  ·  SQLite (zero-server, libSQL adapter)'),
        ('AI Engine',      'Azure AI Foundry  ·  GPT-5.1  via @azure-rest/ai-inference SDK'),
        ('Browser Engine', 'Playwright 1.60  ·  Chromium  ·  used for both crawling and test execution'),
        ('Real-time',      'Server-Sent Events  ·  NestJS @Sse()  ·  RxJS Subjects  ·  no polling'),
        ('Monorepo',       'npm workspaces  ·  apps/api (NestJS)  ·  apps/web (Next.js)  ·  packages/shared'),
    ]

    tbl = doc.add_table(rows=len(rows) + 1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    for ci, h in enumerate(['LAYER', 'TECHNOLOGY']):
        cell = tbl.rows[0].cells[ci]
        set_cell_bg(cell, '7500C0')
        p = cell.paragraphs[0]
        para_spacing(p, 50, 50)
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(8.5)
        r.font.color.rgb = WHITE
        r.font.name = 'Arial'

    for i, (layer, tech) in enumerate(rows):
        row = tbl.rows[i + 1]
        bg = 'F4F4F4' if i % 2 == 0 else 'FFFFFF'
        set_cell_bg(row.cells[0], '111111')
        set_cell_bg(row.cells[1], bg)
        row.cells[0].width = Inches(1.1)
        row.cells[1].width = Inches(5.5)

        lp = row.cells[0].paragraphs[0]
        tp = row.cells[1].paragraphs[0]
        para_spacing(lp, 45, 45)
        para_spacing(tp, 45, 45)

        lr = lp.add_run(layer)
        lr.bold = True
        lr.font.size = Pt(9)
        lr.font.color.rgb = PURPLE
        lr.font.name = 'Arial'

        tr = tp.add_run(tech)
        tr.font.size = Pt(9)
        tr.font.name = 'Arial'


# ── Section 4: Business Impact ─────────────────────────────────────────────────

def make_impact(doc):
    section_label(doc, '04  Business Impact')

    kpis = [
        ('Test Setup\nTime',    '16 hrs', '0 hrs',    '−100%'),
        ('Test\nCoverage',      '40%',    '90%',       '+125%'),
        ('Defect\nDetection',   '55%',    '88%',       '+60%'),
        ('Manual QA\nEffort',   '24 hrs', '<2 hrs',    '−92%'),
        ('Time to\nReport',     '60 min', '<3 min',    '−95%'),
    ]

    tbl = doc.add_table(rows=1, cols=len(kpis))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    chip_w = 6.6 / len(kpis)
    for ci, (metric, before, after, delta) in enumerate(kpis):
        cell = tbl.rows[0].cells[ci]
        cell.width = Inches(chip_w)
        set_cell_bg(cell, '0d0d0d')

        mp = cell.paragraphs[0]
        mp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para_spacing(mp, 60, 20)
        mr = mp.add_run(metric)
        mr.bold = True
        mr.font.size = Pt(8.5)
        mr.font.color.rgb = PURPLE
        mr.font.name = 'Arial'

        bp = cell.add_paragraph()
        bp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para_spacing(bp, 10, 5)
        br = bp.add_run(before)
        br.font.size = Pt(9)
        br.font.color.rgb = RGBColor(0xFF, 0x55, 0x55)
        br.font.name = 'Arial'

        ap = cell.add_paragraph()
        ap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para_spacing(ap, 5, 5)
        ar = ap.add_run(after)
        ar.font.size = Pt(9.5)
        ar.bold = True
        ar.font.color.rgb = RGBColor(0x00, 0xCC, 0x66)
        ar.font.name = 'Arial'

        dp = cell.add_paragraph()
        dp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para_spacing(dp, 8, 60)
        dr = dp.add_run(delta)
        dr.bold = True
        dr.font.size = Pt(11)
        dr.font.color.rgb = RGBColor(0x00, 0xCC, 0x66)
        dr.font.name = 'Arial'


# ── Section 5: Scalability ─────────────────────────────────────────────────────

def make_scalability(doc):
    section_label(doc, '05  Scalability Roadmap')

    cols = [
        ('NEAR-TERM  (1–2 sprints)', [
            'CI/CD integration — GitHub Actions / Jenkins plugin',
            'BullMQ job queue — concurrent multi-user execution',
            'Self-healing selectors — GPT-5.1 auto-repairs broken CSS selectors',
        ]),
        ('MEDIUM-TERM  (1 quarter)', [
            'Natural language test commands ("Test the checkout flow")',
            'Multi-browser execution — Firefox, WebKit',
            'Scheduled runs with email / Slack notifications',
        ]),
        ('LONG-TERM  (scale)', [
            'Accenture myWizard integration — surface reports in delivery dashboards',
            'Cross-project trend analysis and quality benchmarking',
            'Export reports as PDF / HTML for client-facing QA gates',
        ]),
    ]

    tbl = doc.add_table(rows=1, cols=3)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    col_w = 6.6 / 3
    for ci, (header, items) in enumerate(cols):
        cell = tbl.rows[0].cells[ci]
        cell.width = Inches(col_w)
        bg = '111111' if ci % 2 == 0 else '0d0d0d'
        set_cell_bg(cell, bg)

        hp = cell.paragraphs[0]
        para_spacing(hp, 60, 30)
        hr = hp.add_run(header)
        hr.bold = True
        hr.font.size = Pt(8.5)
        hr.font.color.rgb = PURPLE
        hr.font.name = 'Arial'

        for item in items:
            ip = cell.add_paragraph()
            para_spacing(ip, 20, 20)
            ir = ip.add_run(f'•  {item}')
            ir.font.size = Pt(9)
            ir.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
            ir.font.name = 'Arial'

        pp = cell.add_paragraph()
        para_spacing(pp, 40, 40)


# ── Footer ─────────────────────────────────────────────────────────────────────

def make_footer(doc):
    add_hr(doc, 'A100FF', '6')
    p = doc.add_paragraph()
    para_spacing(p, 40, 0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        'TESTA  ·  AI-Powered Autonomous Web Testing Platform  ·  '
        'Accenture DCX Innovation  ·  INNOVATE X  ·  May 2026'
    )
    r.font.size = Pt(8)
    r.font.color.rgb = GREY_DARK
    r.font.name = 'Arial'


# ── Entry point ────────────────────────────────────────────────────────────────

def build_summary():
    doc = Document()
    for section in doc.sections:
        section.top_margin    = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin   = Cm(1.8)
        section.right_margin  = Cm(1.8)

    make_header(doc)
    make_problem(doc)
    make_solution(doc)
    make_tech_stack(doc)
    make_impact(doc)
    make_scalability(doc)
    make_footer(doc)
    return doc


if __name__ == '__main__':
    import os
    out = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'TESTA_Solution_Summary.docx')
    print('Building solution summary ...')
    doc = build_summary()
    doc.save(out)
    print(f'Saved: {out}')
